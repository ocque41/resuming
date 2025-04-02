import os
import io
import logging
import boto3
import json
from typing import Dict, Optional, Any, Tuple
from botocore.exceptions import ClientError
import fitz  # PyMuPDF for PDF handling
import docx  # python-docx for DOCX handling

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize S3 client
s3_client = boto3.client(
    's3',
    aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
    aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
    region_name=os.getenv('AWS_REGION', 'us-east-1')
)

# Get S3 bucket name from environment variable
S3_BUCKET_NAME = os.getenv('S3_BUCKET_NAME')
if not S3_BUCKET_NAME:
    logger.warning("S3_BUCKET_NAME environment variable not set. Document retrieval will fail.")

class DocumentHandler:
    """Handler for document operations including retrieval and processing."""
    
    @staticmethod
    async def get_document_from_s3(s3_key: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """
        Retrieve a document from S3 and extract its text content.
        
        Args:
            s3_key: The S3 key of the document
            
        Returns:
            Tuple containing:
            - Document text content (or None if failed)
            - Document metadata (or None if failed)
        """
        if not S3_BUCKET_NAME:
            logger.error("S3_BUCKET_NAME not configured")
            return None, None
            
        try:
            # Get file from S3
            response = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
            content_type = response.get('ContentType', '')
            file_content = response['Body'].read()
            
            # Extract text based on file type
            text_content = None
            
            if content_type == 'application/pdf' or s3_key.lower().endswith('.pdf'):
                text_content = DocumentHandler.extract_text_from_pdf(file_content)
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or s3_key.lower().endswith('.docx'):
                text_content = DocumentHandler.extract_text_from_docx(file_content)
            elif content_type == 'text/plain' or s3_key.lower().endswith('.txt'):
                text_content = file_content.decode('utf-8')
            elif content_type == 'application/json' or s3_key.lower().endswith('.json'):
                json_content = json.loads(file_content.decode('utf-8'))
                # If it's a JSON document, check if it has a 'content' field
                if isinstance(json_content, dict) and 'content' in json_content:
                    text_content = json_content['content']
                else:
                    text_content = json.dumps(json_content, indent=2)
            else:
                logger.warning(f"Unsupported content type: {content_type} for key: {s3_key}")
                # Attempt to decode as text as fallback
                try:
                    text_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    logger.error(f"Could not decode file as text: {s3_key}")
                    return None, None
            
            # Extract basic metadata
            metadata = {
                'content_type': content_type,
                'size': response.get('ContentLength', 0),
                'last_modified': response.get('LastModified', '').isoformat() if response.get('LastModified') else None,
                's3_key': s3_key,
                'filename': s3_key.split('/')[-1] if '/' in s3_key else s3_key
            }
            
            return text_content, metadata
            
        except ClientError as e:
            logger.error(f"Error retrieving document from S3: {str(e)}")
            return None, None
        except Exception as e:
            logger.error(f"Unexpected error processing document: {str(e)}")
            return None, None
    
    @staticmethod
    def extract_text_from_pdf(pdf_content: bytes) -> str:
        """Extract text content from a PDF file."""
        try:
            # Create a file-like object from the bytes
            pdf_file = io.BytesIO(pdf_content)
            
            # Open the PDF with PyMuPDF
            doc = fitz.open(stream=pdf_file, filetype="pdf")
            
            # Extract text from each page
            text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text() + "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return "Error extracting text from PDF document"
    
    @staticmethod
    def extract_text_from_docx(docx_content: bytes) -> str:
        """Extract text content from a DOCX file."""
        try:
            # Create a file-like object from the bytes
            docx_file = io.BytesIO(docx_content)
            
            # Open the DOCX with python-docx
            doc = docx.Document(docx_file)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs]
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n\n".join(paragraphs) + "\n\n" + "\n".join(tables_text)
            
            return all_text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return "Error extracting text from DOCX document"
            
    @staticmethod
    async def get_document_from_database(document_id: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """
        Retrieve document content and metadata from database by document ID.
        This will be implemented when we have a database connection.
        
        Args:
            document_id: The document ID in the database
            
        Returns:
            Tuple containing:
            - Document text content (or None if failed)
            - Document metadata (or None if failed)
        """
        # This is a placeholder - will be implemented with actual database
        logger.info(f"Requested document with ID: {document_id}")
        
        # For now, return a sample document for testing
        if document_id == "test-document":
            return (
                "This is a sample document for testing purposes.\n\nIt contains multiple paragraphs and can be used to test the AI agent's document handling capabilities.",
                {"id": document_id, "name": "Test Document", "type": "text"}
            )
        
        return None, None 